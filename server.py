"""
Word 格式修改器 — 本地 Web 服务
提供 API 接口供 HTML 界面调用。
依赖：flask （pip install flask）
"""
import json
import os
import socket
import uuid
import threading
from pathlib import Path
from flask import Flask, request, jsonify, send_file, send_from_directory
from werkzeug.utils import secure_filename

from formatter import scan_document, format_document, clean_text


def is_port_in_use(port: int) -> bool:
    """检测指定端口是否已被占用"""
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        return s.connect_ex(("127.0.0.1", port)) == 0

app = Flask(__name__)
BASE_DIR   = Path(__file__).parent
CONFIG_PATH = BASE_DIR / "config.json"
UPLOAD_DIR = BASE_DIR / "uploads"
RESULT_DIR  = BASE_DIR / "results"

UPLOAD_DIR.mkdir(exist_ok=True)
RESULT_DIR.mkdir(exist_ok=True)

# ── 工具函数 ──────────────────────────────────────────────

def load_config():
    return json.loads(CONFIG_PATH.read_text("utf-8"))


def save_config(data):
    CONFIG_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), "utf-8")


def get_preset(name: str):
    cfg = load_config()
    return cfg["presets"].get(name)


def find_uploaded_file(file_id: str) -> Path | None:
    files = list(UPLOAD_DIR.glob(f"{file_id}_*"))
    return files[0] if files else None


# ── .txt / .md 转换 ───────────────────────────────────────

def convert_txt_to_docx(txt_path: Path, out_path: Path):
    """将 .txt 文件转换为 .docx（每段一行）"""
    from docx import Document
    doc = Document()
    text = txt_path.read_text("utf-8")
    for line in text.splitlines():
        line = line.rstrip()
        if line:
            doc.add_paragraph(line)
        else:
            # 空行也加一个空段落，保持结构
            doc.add_paragraph("")
    doc.save(str(out_path))


def convert_md_to_docx(md_path: Path, out_path: Path):
    """将 .md 文件转换为 .docx（解析 # ## ### 标题）"""
    from docx import Document
    doc = Document()
    text = md_path.read_text("utf-8")
    lines = text.splitlines()

    in_code_block = False
    in_table = False

    for line in lines:
        stripped = line.strip()

        # 跳过代码块
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            doc.add_paragraph(line)
            continue

        # 标题行
        if stripped.startswith("# "):
            doc.add_paragraph(stripped[2:])
        elif stripped.startswith("## "):
            doc.add_paragraph(stripped[3:])
        elif stripped.startswith("### "):
            doc.add_paragraph(stripped[4:])
        elif stripped.startswith("#### "):
            doc.add_paragraph(stripped[5:])
        elif stripped == "":
            # 空行，跳过（不添加空段落）
            continue
        else:
            doc.add_paragraph(line)

    doc.save(str(out_path))


SUPPORTED_EXTENSIONS = {".docx", ".txt", ".md"}


# ── 路由 ──────────────────────────────────────────────────

@app.route("/")
def index():
    html_path = BASE_DIR / "index.html"
    if html_path.exists():
        return html_path.read_text("utf-8")
    return "index.html not found", 404


@app.route("/api/config")
def api_config():
    return jsonify(load_config())


@app.route("/api/presets", methods=["GET"])
def api_presets():
    cfg = load_config()
    presets = [
        {"name": n, "description": p.get("description", "")}
        for n, p in cfg["presets"].items()
    ]
    return jsonify({"presets": presets})


@app.route("/api/preset/<name>")
def api_preset(name: str):
    preset = get_preset(name)
    if preset is None:
        return jsonify({"error": "预设不存在"}), 404
    return jsonify(preset)


@app.route("/api/preset/<name>", methods=["PUT"])
def api_save_preset(name: str):
    data = request.get_json()
    cfg = load_config()
    cfg["presets"][name] = data
    save_config(cfg)
    return jsonify({"ok": True, "message": f"预设「{name}」已保存"})


@app.route("/api/preset/<name>", methods=["DELETE"])
def api_delete_preset(name: str):
    cfg = load_config()
    if name in cfg["presets"]:
        del cfg["presets"][name]
        save_config(cfg)
        return jsonify({"ok": True})
    return jsonify({"error": "预设不存在"}), 404


# ── 字体扫描 ─────────────────────────────────────────────

def _get_font_display_name(font_path: Path) -> str | None:
    """
    用 fonttools 读取字体的显示名称（优先中文）。
    返回字体全名（nameID=4），如果读取失败返回 None。
    """
    try:
        from fontTools.ttLib import TTFont
    except ImportError:
        return None

    try:
        font = TTFont(str(font_path))
        name_table = font["name"]

        # 优先：中文简体 (platformID=3, langID=0x804)
        for record in name_table.names:
            if record.nameID == 4 and record.platformID == 3 and record.langID == 0x804:
                try:
                    return record.toUnicode()
                except Exception:
                    pass

        # 次选：中文 (platformID=1, langID=0x21)
        for record in name_table.names:
            if record.nameID == 4 and record.platformID == 1 and record.langID == 0x21:
                try:
                    return record.toUnicode()
                except Exception:
                    pass

        # fallback：任何 nameID=4
        for record in name_table.names:
            if record.nameID == 4:
                try:
                    val = record.toUnicode()
                    if val:
                        return val
                except Exception:
                    pass
    except Exception:
        pass
    return None


def scan_system_fonts() -> list:
    """
    扫描系统上安装的字体，返回 [{"name", "path"}, ...]。
    使用 fonttools 读取字体内部名称（优先中文）。
    递归扫描字体目录（含子目录）。
    """
    fonts = []
    seen  = set()

    import platform
    system = platform.system()
    home = Path.home()

    if system == "Darwin":
        dirs = [
            Path("/System/Library/Fonts"),
            Path("/Library/Fonts"),
            home / "Library/Fonts",
        ]
    elif system == "Windows":
        dirs = [
            Path("C:/Windows/Fonts"),
            home / "AppData/Local/Microsoft/Windows/Fonts",
        ]
    else:
        dirs = [
            Path("/usr/share/fonts"),
            Path("/usr/local/share/fonts"),
            home / ".fonts",
        ]

    for d in dirs:
        if not d.exists():
            continue
        # rglob 递归扫描所有子目录
        for ext in ("*.ttf", "*.ttc", "*.otf", "*.TTF", "*.TTC", "*.OTF"):
            for f in d.rglob(ext):
                display_name = _get_font_display_name(f)
                if not display_name:
                    display_name = f.stem  # fallback 到文件名
                if display_name not in seen:
                    seen.add(display_name)
                    fonts.append({"name": display_name, "path": str(f)})

    return fonts


@app.route("/api/fonts")
def api_fonts():
    """返回系统上安装的中文字体列表"""
    fonts = scan_system_fonts()
    # 常见字体推荐（排在前面）
    recommended = [
        "方正小标宋简体", "FZXiaoBiaoSong", "FZSXBSK—GBK",
        "仿宋 GB2312", "FangSong GB2312", "FZFangSong-Z02S", "方正仿宋简体",
        "楷体 GB2312", "KaiTi GB2312", "STKaiti", "华文楷体",
        "黑体", "SimHei", "STHeiti", "华文黑体",
        "宋体", "SimSun", "STSong", "华文宋体",
        "微软雅黑", "Microsoft YaHei",
        "PingFang SC", "苹方",
    ]
    rec_set = set(recommended)
    ordered = [f for f in fonts if f["name"] in rec_set]
    others  = [f for f in fonts if f["name"] not in rec_set]
    return jsonify({"ok": True, "recommended": ordered, "all": others})


# ── 文件上传 ──────────────────────────────────────────────

@app.route("/api/upload", methods=["POST"])
def api_upload():
    if "files" not in request.files:
        return jsonify({"error": "没有文件"}), 400

    file_ids = []
    for file in request.files.getlist("files"):
        if file.filename == "":
            continue

        file_id   = str(uuid.uuid4())
        orig_name = file.filename
        ext       = Path(orig_name).suffix.lower()

        if ext not in SUPPORTED_EXTENSIONS:
            continue

        # 保存上传文件
        safe_name  = secure_filename(orig_name)
        raw_path   = UPLOAD_DIR / f"{file_id}_raw{safe_name}"
        file.save(str(raw_path))

        final_path = raw_path
        final_name = safe_name

        # .txt / .md → 转 .docx
        if ext in (".txt", ".md"):
            docx_name = Path(safe_name).stem + ".docx"
            docx_path = UPLOAD_DIR / f"{file_id}_{docx_name}"
            try:
                if ext == ".txt":
                    convert_txt_to_docx(raw_path, docx_path)
                else:
                    convert_md_to_docx(raw_path, docx_path)
                raw_path.unlink(missing_ok=True)
                final_path = docx_path
                final_name = docx_name
            except Exception as e:
                raw_path.unlink(missing_ok=True)
                file_ids.append({
                    "file_id":      file_id,
                    "original_name": orig_name,
                    "error":         f"转换失败：{e}",
                })
                continue

        file_ids.append({
            "file_id":      file_id,
            "filename":      final_name,
            "original_name": orig_name,
        })

        # 保存原文件名到 meta 文件（用于输出文件名）
        meta = {"original_name": orig_name}
        meta_path = UPLOAD_DIR / f"{file_id}.meta.json"
        meta_path.write_text(json.dumps(meta, ensure_ascii=False), encoding="utf-8")

    return jsonify({"ok": True, "files": file_ids})


# ── 扫描（预览） ─────────────────────────────────────────

@app.route("/api/font-mapping", methods=["GET"])
def api_get_font_mapping():
    """返回当前字体映射"""
    cfg = load_config()
    return jsonify({"ok": True, "mapping": cfg.get("font_mapping", {})})


@app.route("/api/font-mapping", methods=["PUT"])
def api_save_font_mapping():
    """保存字体映射"""
    data = request.get_json()
    cfg = load_config()
    cfg["font_mapping"] = data.get("mapping", {})
    save_config(cfg)
    return jsonify({"ok": True, "message": "字体映射已保存"})


@app.route("/api/test-font", methods=["POST"])
def api_test_font():
    """用指定字体名生成一个测试 .docx，返回下载链接"""
    data = request.get_json()
    font_cn = data.get("font_cn", "")
    font_en = data.get("font_en", font_cn)
    from docx import Document
    from docx.shared import Pt
    from formatter import _set_run_font
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run(f"字体测试：{font_cn} / {font_en}")
    _set_run_font(run, font_cn, font_en, 17, False)
    out = RESULT_DIR / "font_test.docx"
    doc.save(str(out))
    return jsonify({"ok": True, "download_url": f"/api/download?path={out}"})

@app.route("/api/scan/<file_id>")
def api_scan(file_id: str):
    """
    扫描文件，返回标题识别结果。
    查询参数：title_index, author_index（默认 0, 1）
    """
    f = find_uploaded_file(file_id)
    if not f:
        return jsonify({"error": "文件不存在"}), 404

    try:
        ti = int(request.args.get("title_index", 0))
        ai = int(request.args.get("author_index", 1))
        result = scan_document(str(f), title_index=ti, author_index=ai)
        return jsonify({"ok": True, **result})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500



# ── 格式化（核心） ───────────────────────────────────────

@app.route("/api/format", methods=["POST"])
def api_format():
    """
    请求体：
      file_id, preset_name,
      title_index, author_index,
      author_position ("before" / "after"),
      output_dir (可选，默认同目录),
      levels: {段落索引: 级别}  （用户手动调整后的）
    """
    data = request.get_json()
    file_id  = data.get("file_id")
    pname   = data.get("preset_name", "党政公文标准")
    ti      = data.get("title_index", 0)
    ai      = data.get("author_index", 1)
    apos    = data.get("author_position", "before")
    out_dir = data.get("output_dir", "")   # 空 = 同目录（原文件旁）

    f = find_uploaded_file(file_id)
    if not f:
        return jsonify({"error": "文件不存在"}), 404

    cfg = load_config()
    preset = cfg["presets"].get(pname)
    if preset is None:
        return jsonify({"error": f"预设「{pname}」不存在"}), 404

    font_mapping = cfg.get("font_mapping", {})

    input_path  = str(f)

    # 从 meta 文件读取原文件名（用于生成输出文件名）
    meta_path = UPLOAD_DIR / f"{file_id}.meta.json"
    if meta_path.exists():
        meta = json.loads(meta_path.read_text("utf-8"))
        original_name = meta["original_name"]
    else:
        # 降级：从文件名推断（兼容旧版未生成 meta 的文件）
        original_name = f.name[len(file_id) + 1:]
        if original_name.startswith("raw"):
            original_name = original_name[3:]

    # 根据原文件扩展名生成输出文件名
    orig_stem = Path(original_name).stem
    orig_ext  = Path(original_name).suffix.lower()
    if orig_ext == ".docx":
        output_name = f"{orig_stem}Format.docx"
    elif orig_ext == ".txt":
        output_name = f"{orig_stem}TXTFormat.docx"
    elif orig_ext == ".md":
        output_name = f"{orig_stem}MDFormat.docx"
    else:
        output_name = f"{orig_stem}Format.docx"

    # 确定输出路径
    if out_dir and Path(out_dir).exists():
        output_path = str(Path(out_dir) / output_name)
    else:
        # 默认：与原文件同目录
        output_path = str(f.parent / output_name)

    # 转换 levels（JSON key 是字符串）
    user_levels = None
    if "levels" in data:
        user_levels = {int(k): v for k, v in data["levels"].items()}

    try:
        result = format_document(
            input_path, output_path, preset,
            user_levels  = user_levels,
            title_index  = ti,
            author_index = ai,
            author_position = apos,
            font_mapping  = font_mapping,
        )
        # 返回相对 URL 和绝对路径
        download_url = f"/api/download?path={output_path}"
        return jsonify({
            "ok": True,
            "output_path":  output_path,
            "output_name":  output_name,
            "download_url":  download_url,
            "saved_to":      output_path,
            **result
        })
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


# ── 下载 ──────────────────────────────────────────────────

@app.route("/api/download")
def api_download():
    """
    通过查询参数 ?path=绝对路径 或 ?name=文件名（results/ 目录）下载
    """
    path = request.args.get("path", "")
    name = request.args.get("name", "")

    if path and Path(path).exists():
        p = Path(path)
        return send_file(str(p), as_attachment=True, download_name=p.name)

    if name:
        p = RESULT_DIR / name
        if p.exists():
            return send_file(str(p), as_attachment=True, download_name=name)

    return jsonify({"error": "文件不存在"}), 404


# ── 状态 ──────────────────────────────────────────────────

@app.route("/api/status")
def api_status():
    return jsonify({
        "status":   "running",
        "uploaded":  len(list(UPLOAD_DIR.glob("*"))),
        "formatted": len(list(RESULT_DIR.glob("*"))),
    })


# ── 启动 ──────────────────────────────────────────────────

if __name__ == "__main__":
    import webbrowser

    if is_port_in_use(8080):
        print("=" * 50)
        print("  ⚠️  端口 8080 已被占用！")
        print("  可能已有 Word 格式修改器在运行。")
        print("  请先关闭之前的终端窗口，再重新启动。")
        print("=" * 50)
        exit(1)

    def open_browser():
        webbrowser.open("http://localhost:8080")

    print("=" * 50)
    print("  Word 格式修改器 启动中...")
    print("  浏览器将自动打开：http://localhost:8080")
    print("  关闭此窗口即可退出")
    print("=" * 50)

    threading.Timer(1.0, open_browser).start()
    app.run(host="127.0.0.1", port=8080, debug=False)
