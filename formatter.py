"""
Word 格式修改器 — 核心处理模块
依赖：python-docx （pip install python-docx）
"""
import re
import copy
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree

# ── 编号格式映射 ─────────────────────────────────────────────
# chinese_dash  → 一、二、三、
# paren_chinese → （一）（二）（三）
# dot_number  → 1. 2. 3.
# dot_dot_number → 1.1 1.2 2.1
# dot_dot_dot_number → 1.1.1 1.1.2
# circle_number → ① ② ③

CHINESE_NUMBERS = "一二三四五六七八九十"

def chinese_number(n: int) -> str:
    """将数字转为中文数字（支持 1-99）"""
    if n <= 10:
        return CHINESE_NUMBERS[n - 1]
    tens = n // 10
    ones = n % 10
    if ones == 0:
        return CHINESE_NUMBERS[tens - 1] + "十"
    if tens == 1:
        return "十" + CHINESE_NUMBERS[ones - 1]
    return CHINESE_NUMBERS[tens - 1] + "十" + CHINESE_NUMBERS[ones - 1]


def format_number(level: int, counter: int, parent_counters: dict) -> str:
    """根据级别和计数器生成编号文字"""
    if level == 1:
        return f"{chinese_number(counter)}、"
    elif level == 2:
        return f"（{chinese_number(counter)}）"
    elif level == 3:
        return f"{counter}. "
    elif level == 4:
        return f"（{counter}）"
    return ""


# ── 标题识别：四级策略 ──────────────────────────────────────────

# 策略二：文字规律正则
PATTERNS = [
    (re.compile(r'^[一二三四五六七八九十]{1,3}、[　\s]?'), 1, "ch_dash"),
    (re.compile(r'^（[一二三四五六七八九十]{1,3}）[　\s]?'), 2, "paren_ch"),
    (re.compile(r'^\d{1,2}\.[　\s]?'), 3, "dot"),
    (re.compile(r'^\d{1,2}\.\d{1,2}[　\s]?'), 3, "dot_dot"),
    (re.compile(r'^[①②③④⑤⑥⑦⑧⑨⑩][　\s]?'), 4, "circle"),
]


def detect_title_level(text: str):
    """策略二：用正则匹配段落开头，返回级别 1-4，无匹配返回 None"""
    stripped = text.strip()
    for pattern, level, _ in PATTERNS:
        if pattern.search(stripped):
            return level
    return None


def infer_levels(paragraph_texts: list, title_index: int = 0, author_index: int = 1):
    """
    扫描全文，推断标题级别。
    返回 (levels_dict, title_idx, author_idx)
      levels_dict: {段落索引: 级别1-4, 未识别=0}
      title_idx:    文章标题段落索引（-1 表示未识别）
      author_idx:   作者段落索引（-1 表示未识别）
    """
    levels = {}
    found_formats = []

    for idx, text in enumerate(paragraph_texts):
        stripped = text.strip()
        matched = False
        for pattern, level, fmt in PATTERNS:
            if pattern.search(stripped):
                levels[idx] = level
                found_formats.append((idx, level, fmt))
                matched = True
                break
        if not matched:
            levels[idx] = 0  # 正文

    # 多种格式时，按首次出现顺序确定层级
    unique_fmts = []
    seen = set()
    for _, _, fmt in found_formats:
        if fmt not in seen:
            unique_fmts.append(fmt)
            seen.add(fmt)

    if len(unique_fmts) >= 2:
        fmt_to_level = {fmt: i + 1 for i, fmt in enumerate(unique_fmts)}
        for idx, level, fmt in found_formats:
            levels[idx] = fmt_to_level.get(fmt, level)
    elif len(found_formats) > 0 and len(unique_fmts) == 1:
        # 只有一种编号格式，全部标为一级
        for idx, _, _ in found_formats:
            levels[idx] = 1

    # 标记文章标题和作者
    if title_index >= 0 and title_index < len(paragraph_texts):
        levels[title_index] = -1
    if author_index >= 0 and author_index < len(paragraph_texts) and author_index != title_index:
        levels[author_index] = -2

    return levels, title_index, author_index


def scan_document(input_path: str, title_index: int = 0, author_index: int = 1):
    """
    扫描文档，返回识别结果（不修改文件）。
    用于预览确认页。
    """
    doc = Document(input_path)
    para_texts = [p.text for p in doc.paragraphs]
    levels, ti, ai = infer_levels(para_texts, title_index, author_index)

    results = []
    for idx, text in enumerate(para_texts):
        level = levels.get(idx, 0)
        if idx == ti:
            label = "文章标题"
        elif idx == ai:
            label = "作者"
        else:
            label = {0: "正文", 1: "一级", 2: "二级", 3: "三级", 4: "四级"}.get(level, "正文")
        results.append({
            "index": idx,
            "text": text[:80] + ("..." if len(text) > 80 else ""),
            "level": level,
            "level_label": label,
            "is_title": idx == ti,
            "is_author": idx == ai,
        })
    return {"total": len(para_texts), "results": results,
            "title_index": ti, "author_index": ai}


# ── 清洗 ───────────────────────────────────────────────────────

MARKDOWN_RE = [
    (re.compile(r'^#+\s*'), ''),
    (re.compile(r'[*_]{2,}'), ''),
    (re.compile(r'[*_]{1}'), ''),
    (re.compile(r'^>\s*'), ''),
    (re.compile(r'^[-*+]\s+'), ''),
    (re.compile(r'\[([^\]]+)\]\([^\)]+\)'), r'\1'),
]


def clean_text(text: str) -> str:
    """清洗 Markdown 残留符号和多余空格"""
    result = text
    for pattern, repl in MARKDOWN_RE:
        result = pattern.sub(repl, result)
    result = re.sub(r'[ 　]{3,}', ' ', result)
    result = result.strip()
    return result


def clean_paragraphs(doc: Document):
    """清洗文档中所有段落的 Markdown 残留"""
    for para in doc.paragraphs:
        cleaned = clean_text(para.text)
        if para.text != cleaned:
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = cleaned
            else:
                para.add_run(cleaned)


# ── 重新编号 ───────────────────────────────────────────────────

def renumber_paragraphs(paragraph_texts: list, levels: dict,
                        title_index: int = -1, author_index: int = -1,
                        style: str = "standard") -> list:
    """
    根据识别到的级别，重新生成编号文字。
    标题和作者不编号。
    """
    counters = {1: 0, 2: 0, 3: 0, 4: 0}
    new_texts = []

    for idx, text in enumerate(paragraph_texts):
        if idx == title_index or idx == author_index:
            new_texts.append(text)  # 标题/作者不改编号
            continue
        level = levels.get(idx, 0)
        if level == 0:
            new_texts.append(text)
            continue

        counters[level] += 1
        for l in range(level + 1, 5):
            counters[l] = 0

        if style == "standard":
            if level == 1:
                new_num = f"{chinese_number(counters[1])}、"
            elif level == 2:
                new_num = f"（{chinese_number(counters[2])}）"
            elif level == 3:
                new_num = f"{counters[3]}. "
            elif level == 4:
                new_num = f"（{counters[4]}）"
            else:
                new_num = ""
        else:  # academic
            if level == 1:
                new_num = f"{counters[1]}. "
            elif level == 2:
                new_num = f"{counters[1]}.{counters[2]}. "
            elif level == 3:
                new_num = f"{counters[1]}.{counters[2]}.{counters[3]}. "
            else:
                new_num = ""

        old_text = strip_number_prefix(text)
        new_texts.append(new_num + old_text)

    return new_texts


def strip_number_prefix(text: str) -> str:
    """去掉段落开头的编号前缀，保留后面的文字"""
    stripped = text.strip()
    for pattern, _, _ in PATTERNS:
        m = pattern.match(stripped)
        if m:
            return stripped[m.end():].strip()
    return stripped


# ── 样式套用 ──────────────────────────────────────────────────

def _set_run_font(run, font_cn: str, font_en: str, size_pt: float, bold: bool = False):
    """设置 run 的字体（含东亚文字字体），并清除主题字体引用"""
    run.font.name = font_en
    run.font.size = Pt(size_pt)
    run.font.bold = bold

    rPr = run._r.get_or_add_rPr()

    # 清除主题字体引用（这些会覆盖手动设置的字体，是Word不生效的根因）
    for theme_attr in ('w:themeEastAsia', 'w:themeAscii', 'w:themeHAnsi', 'w:themeCs'):
        el = rPr.find(qn(theme_attr))
        if el is not None:
            rPr.remove(el)

    # 设置 rFonts 元素
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))

    # 中文字符渲染用字体（w:eastAsia）
    rFonts.set(qn('w:eastAsia'), font_cn)
    # 西文字符渲染用字体（w:ascii / w:hAnsi）
    rFonts.set(qn('w:ascii'),    font_en)
    rFonts.set(qn('w:hAnsi'),   font_en)
    # 复杂文字脚本（w:cs）
    rFonts.set(qn('w:cs'),        font_en)

    # 同时设置字号（东亚和西文都要设）
    for sz_tag in ('w:sz', 'w:szCs'):
        sz = rPr.find(qn(sz_tag))
        if sz is None:
            sz = etree.SubElement(rPr, qn(sz_tag))
        sz.set(qn('w:val'), str(int(size_pt * 2)))


def _apply_para_format(para, cfg: dict):
    """应用段落格式（对齐、行距、段间距、首行缩进）"""
    pf = para.paragraph_format
    align = cfg.get("alignment", "left")
    if align == "center":
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    rule = cfg.get("line_spacing_rule", "multiple")
    if rule == "fixed":
        pf.line_spacing = Pt(cfg.get("line_spacing_pt", 30))
    elif rule == "multiple":
        pf.line_spacing = cfg.get("line_spacing_multiple", 1.34)

    pf.space_before = Pt(cfg.get("space_before", 0))
    pf.space_after = Pt(cfg.get("space_after", 0))

    indent = cfg.get("first_line_indent", 0)
    if indent > 0:
        pf.first_line_indent = Pt(indent * 5.5)


def _apply_font_mapping(cfg: dict, font_mapping: dict | None) -> dict:
    """
    把 font_mapping 应用到 cfg 的所有 font_cn / font_en 值上。
    返回一份新的 cfg（不修改原 dict）。
    """
    if not font_mapping:
        return cfg
    import copy
    cfg = copy.deepcopy(cfg)
    for key, val in cfg.items():
        if isinstance(val, dict):
            if "font_cn" in val and val["font_cn"] in font_mapping:
                val["font_cn"] = font_mapping[val["font_cn"]]
            if "font_en" in val and val["font_en"] in font_mapping:
                val["font_en"] = font_mapping[val["font_en"]]
    # page_number 里的字体
    pn = cfg.get("page_number", {})
    if "font_cn" in pn and pn["font_cn"] in font_mapping:
        pn["font_cn"] = font_mapping[pn["font_cn"]]
    if "font_en" in pn and pn["font_en"] in font_mapping:
        pn["font_en"] = font_mapping[pn["font_en"]]
    return cfg


def format_document(input_path: str, output_path: str, preset: dict,
                   user_levels: Optional[dict] = None,
                   title_index: int = 0, author_index: int = 1,
                   author_position: str = "before",
                   font_mapping: dict | None = None):
    """
    主函数：格式化一个 Word 文档。
    font_mapping: {原始字体名: 系统实际字体名} 替换表
    """
    doc = Document(input_path)
    cfg = _apply_font_mapping(preset, font_mapping)

    # 1. 清洗
    clean_paragraphs(doc)

    # 2. 识别标题级别
    para_texts = [p.text for p in doc.paragraphs]
    levels, ti, ai = infer_levels(para_texts, title_index, author_index)
    if user_levels is not None:
        levels = user_levels
        ti = title_index
        ai = author_index

    # 3. 重新编号
    numbering_style = cfg.get("numbering", {}).get("style", "standard")
    new_texts = renumber_paragraphs(
        para_texts, levels, ti, ai, style=numbering_style
    )

    # 4. 按顺序重建文档段落（处理 author_position）
    #    如果 author_position == "after"，把作者段移到文末
    author_para = None
    if author_position == "after" and ai >= 0 and ai < len(doc.paragraphs):
        author_para = doc.paragraphs[ai]
        # 标记，稍后处理
        author_para._author_move = True

    # 5. 逐段套用格式
    for idx, para in enumerate(doc.paragraphs):
        # 跳过将要移动的作者段（会在最后处理）
        if hasattr(para, '_author_move'):
            continue

        new_text = new_texts[idx] if idx < len(new_texts) else para.text

        # 判断此段是什么
        if idx == ti:
            # 文章标题
            fcfg = cfg.get("title", {})
            _format_para(para, new_text, fcfg, is_title=True)
        elif idx == ai:
            # 作者（author_position=before 时）
            fcfg = cfg.get("author", {})
            _format_para(para, new_text, fcfg)
        else:
            level = levels.get(idx, 0)
            if level == 0:
                fcfg = cfg.get("body", {})
            else:
                fcfg = cfg.get(f"heading{level}", cfg.get("body", {}))
            _format_para(para, new_text, fcfg)

    # 如果 author_position=after，在文末追加作者段
    if author_position == "after" and author_para is not None:
        new_para = doc.add_paragraph()
        for run in author_para.runs:
            new_run = new_para.add_run(run.text)
            # 复制字体
            if run.font.name:
                new_run.font.name = run.font.name
            if run.font.size:
                new_run.font.size = run.font.size
            new_run.font.bold = run.font.bold
        fcfg = cfg.get("author", {})
        _apply_para_format(new_para, fcfg)

    # 6. 设置页面
    page_cfg = cfg.get("page", {})
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(page_cfg.get("margin_top", 3.8))
    section.bottom_margin = Cm(page_cfg.get("margin_bottom", 3.2))
    section.left_margin = Cm(page_cfg.get("margin_left", 2.7))
    section.right_margin = Cm(page_cfg.get("margin_right", 2.7))

    # 7. 页码（OXML 正确插入可更新页码域）
    add_page_numbers(doc, cfg.get("page_number", {}))

    doc.save(output_path)
    return {"total_paragraphs": len(doc.paragraphs), "inferred_levels": levels,
            "title_index": ti, "author_index": ai}


def _format_para(para, new_text: str, fcfg: dict, is_title: bool = False):
    """格式化单个段落"""
    # 写文字：保留第一个 run，清空其余
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ''
        first_run = para.runs[0]
    else:
        first_run = para.add_run(new_text)

    # 设置字体
    _set_run_font(
        first_run,
        fcfg.get("font_cn", "仿宋 GB2312"),
        fcfg.get("font_en", "FangSong"),
        fcfg.get("size_pt", 17),
        fcfg.get("bold", False)
    )

    # 设置段落格式
    _apply_para_format(para, fcfg)


# ── 页码插入（OXML）────────────────────────────────────────────

def add_page_numbers(doc: Document, page_cfg: dict):
    """
    通过 OXML 在页脚插入可自动更新的 PAGE 域。
    格式：- 1 -  （{n} 替换为 Word 页码域）
    """
    section = doc.sections[0]
    footer = section.footer
    if footer is None:
        return
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    # 对齐
    pos = page_cfg.get("position", "right_bottom")
    if "center" in pos:
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif "right" in pos:
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 清空页脚现有内容
    p = footer_para._p
    for child in list(p):
        p.remove(child)

    fmt = page_cfg.get("format", "- {n} -")
    left_text, sep, right_text = fmt.partition("{n}")

    # 构建 OXML：w:p → w:r 子元素
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def _add_run(parent, text: str, font_cfg: dict):
        r = etree.SubElement(parent, qn('w:r'))
        # 字体设置
        rPr = etree.SubElement(r, qn('w:rPr'))
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:eastAsia'), font_cfg.get("font_cn", "宋体"))
        rFonts.set(qn('w:ascii'), font_cfg.get("font_en", "SimSun"))
        rFonts.set(qn('w:hAnsi'), font_cfg.get("font_en", "SimSun"))
        sz = etree.SubElement(rPr, qn('w:sz'))
        sz.set(qn('w:val'), str(int(font_cfg.get("size_pt", 14) * 2)))
        szCs = etree.SubElement(rPr, qn('w:szCs'))
        szCs.set(qn('w:val'), str(int(font_cfg.get("size_pt", 14) * 2)))
        # 文字
        t = etree.SubElement(r, qn('w:t'))
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        return r

    # 左侧文字
    if left_text:
        _add_run(p, left_text, page_cfg)

    # PAGE 域（自动更新页码）
    r_fldBegin = etree.SubElement(p, qn('w:r'))
    fldChar_begin = etree.SubElement(r_fldBegin, qn('w:fldChar'))
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    r_instr = etree.SubElement(p, qn('w:r'))
    instrText = etree.SubElement(r_instr, qn('w:instrText'))
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = " PAGE \\* MERGEFORMAT "

    r_fldEnd = etree.SubElement(p, qn('w:r'))
    fldChar_end = etree.SubElement(r_fldEnd, qn('w:fldChar'))
    fldChar_end.set(qn('w:fldCharType'), 'end')

    # 右侧文字
    if right_text:
        _add_run(p, right_text, page_cfg)


# ── 命令行入口 ─────────────────────────────────────────────────

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("用法：python formatter.py <输入.docx> [输出_formatted.docx]")
        sys.exit(1)
    inp = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else inp.replace(".docx", "_formatted.docx")
    from pathlib import Path
    import json
    cfg_path = Path(__file__).parent / "config.json"
    preset = json.loads(cfg_path.read_text("utf-8"))["presets"]["党政公文标准"]
    result = format_document(inp, out, preset)
    print(f"处理完成：{result}")
